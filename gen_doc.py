# -*- coding: utf-8 -*-
"""生成 SmartMentor 项目技术说明 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 全局中文字体
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal'].font.size = Pt(10.5)
try:
    from docx.oxml.ns import qn
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
except Exception:
    pass

NAVY = RGBColor(0x1a, 0x3a, 0x6b)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.name = '微软雅黑'
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1e, 0x4d, 0x8c)
    return p

def para(text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p

print("helpers ready")

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ============ 封面 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('SmartMentor 智学导师')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('基于大模型的个性化资源生成与学习多智能体系统\n技术实现说明书')
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x2a, 0x5a, 0x9e)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('\n第十五届中国软件杯大赛 · A3 赛题（出题企业：科大讯飞）\n').font.size = Pt(11)
doc.add_page_break()
print("cover done")

# ============ 1. 项目概述 ============
h1('一、项目概述')
para('SmartMentor 智学导师是面向第十五届中国软件杯 A3 赛题开发的"基于大模型的个性化资源生成与学习多智能体系统"。'
     '系统以高校专业课程为切入点，通过对话式交互自动构建学生画像，由多个角色化智能体协同生成个性化、多模态的学习资源，'
     '并规划动态学习路径、提供智能辅导与学习效果评估，实现"因材施教"的数字化落地。')
para('系统采用前后端分离架构：后端基于 Spring Boot 3 提供多智能体编排与大模型适配能力，'
     '前端基于 Vue 3 提供流式对话、Markdown 渲染与资源卡片化展示。AI 能力优先使用科大讯飞星火大模型，'
     '并以 DeepSeek 作为高可用回退。')

# ============ 2. 技术栈 ============
h1('二、整体技术栈')

h2('2.1 后端技术栈')
add_table(
    ['类别', '技术 / 框架', '版本', '用途说明'],
    [
        ['语言/运行时', 'Java', '17', '后端开发语言（LTS）'],
        ['核心框架', 'Spring Boot', '3.x', '应用主框架、自动装配、内嵌容器'],
        ['Web', 'spring-boot-starter-web', '3.x', 'RESTful API 与 SSE 流式接口'],
        ['持久层', 'Spring Data JPA / Hibernate', '3.x', 'ORM、知识点掌握度等数据持久化'],
        ['安全', 'Spring Security + JJWT', '—', 'JWT 鉴权、登录态与接口保护'],
        ['缓存', 'Spring Data Redis', '3.x', '学生画像缓存、会话状态'],
        ['数据库', 'MySQL / H2', '8.x / —', '生产用 MySQL，演示用 H2 内存库'],
        ['HTTP 客户端', 'OkHttp', '4.x', '调用星火/DeepSeek 大模型与外部资源'],
        ['邮件', 'spring-boot-starter-mail', '3.x', '验证码 / 通知邮件'],
        ['工具', 'Lombok / hibernate-types', '—', '样板代码消除、JSON 字段映射'],
        ['构建', 'Maven', '3.x', '依赖管理与打包'],
    ],
    widths=[1.2, 2.2, 0.9, 2.5])

print("backend stack done")

h2('2.2 前端技术栈')
add_table(
    ['类别', '技术 / 框架', '版本', '用途说明'],
    [
        ['框架', 'Vue', '3.4', '渐进式前端框架（Composition API）'],
        ['路由', 'Vue Router', '4.3', '单页应用路由'],
        ['构建工具', 'Vite', '5.x', '极速冷启动与热更新、生产打包'],
        ['Markdown', 'marked', '18.x', 'AI 回复的 Markdown 渲染'],
        ['公式', 'KaTeX', '0.16', 'LaTeX 数学公式渲染'],
        ['动画', 'GSAP', '3.15', '页面与交互动画'],
        ['图标', 'RemixIcon', '4.6', '矢量图标库'],
    ],
    widths=[1.2, 2.0, 0.9, 2.7])

h2('2.3 AI 能力与外部服务')
add_table(
    ['类别', '技术 / 服务', '用途说明'],
    [
        ['主用大模型', '科大讯飞星火大模型', '满足赛题"AI 能力优先使用科大讯飞"要求，承担对话与各 Agent 推理'],
        ['回退大模型', 'DeepSeek', '主用未配置或调用失败时自动回退，保证多智能体稳定可用'],
        ['模型协议', 'OpenAI 兼容 Chat Completions', '统一适配层屏蔽不同厂商差异，支持同步与 SSE 流式'],
        ['资源检索', '哔哩哔哩 / 中国大学MOOC / 学堂在线', '学习视频与课程资源检索，失败时降级为多平台搜索直达卡片'],
        ['AI 编程工具', 'Claude Code 等', '开发阶段辅助编码（按赛题要求说明）'],
    ],
    widths=[1.6, 2.4, 3.0])

print("frontend & ai stack done")

# ============ 3. 系统架构 ============
h1('三、系统架构与多智能体设计')
para('系统采用"前端展现层 — 后端服务层 — 多智能体编排层 — 大模型适配层 — 数据存储层"的分层架构。')
h2('3.1 多智能体协同框架')
para('后端实现了事件驱动的多智能体编排器（AgentOrchestrator），通过"事件订阅—触发—级联"机制驱动各角色智能体协作，'
     '并支持线性 Pipeline 与事件级联两种编排模式。为防止级联死循环，设置了最大协作轮次上限。系统包含以下角色化智能体：')
add_table(
    ['智能体', '职责', '产出'],
    [
        ['DiagnosticAgent（诊断）', '分析学生答题与知识掌握，定位薄弱知识点', '诊断会话、薄弱点列表'],
        ['TracingAgent（溯源）', '从错误表现回溯根因知识点', '错因溯源结果'],
        ['PlanningAgent（规划）', '结合画像与知识图谱规划个性化学习路径', '动态学习路径'],
        ['TeachingAgent（教学）', '生成讲解文档、思维导图、练习、实操案例等资源', '多类型学习资源'],
        ['EvaluationAgent（评估）', '多维评估学习效果并动态调整推送策略', '效果评估报告'],
    ],
    widths=[1.8, 3.0, 2.0])
para('编排器与大模型解耦：所有智能体通过统一的 LlmService 适配层调用大模型，'
     '适配层屏蔽星火与 DeepSeek 的协议差异，对上层提供同步、JSON、多轮、流式四类调用能力。')

h2('3.2 个性化资源类型（≥5 种，实际 7 种）')
add_table(
    ['序号', '资源类型', '生成方式'],
    [
        ['1', '讲解文档', 'TeachingAgent 按知识点概念分层生成'],
        ['2', '知识结构思维导图', '中心节点 + 前置基础 + 常见误区结构化输出'],
        ['3', '分层练习题', '按掌握度动态生成不同难度题目'],
        ['4', '拓展阅读', '基于知识点生成阅读方向'],
        ['5', '实操案例', '生成可动手的实践任务/代码案例'],
        ['6', '教学视频推荐', '多平台检索 + 降级搜索直达卡片'],
        ['7', '动画讲解脚本', '问题场景→核心机制→操作步骤→结果反馈'],
    ],
    widths=[0.7, 2.3, 3.5])
print("architecture done")

# ============ 4. 关键功能实现细节 ============
h1('四、关键功能实现细节')

h2('4.1 对话式学习画像（≥6 维度）')
para('画像通过自然语言对话自动抽取特征构建，包含 6 个动态维度：知识基础（knowledgeState）、'
     '易错点偏好（errorPatterns）、学习行为（learningBehavior）、认知风格（cognitiveStyle）、'
     '课程目标（goalProfile）、资源偏好与学习节奏（resourcePreference）。每个维度带置信度评分，'
     '并随学生答题、学习活动数据实时更新。画像结果缓存在 Redis，降低重复计算开销。')

h2('4.2 流式输出（SSE）')
para('对话采用 Server-Sent Events 流式返回，后端通过 SseEmitter 推送 metadata、resources、message、done、error 五类事件；'
     '并以每 15 秒一次的注释行心跳保活，避免反向代理超时断连。前端边接收边渲染，配合打字光标动画呈现"边想边答"效果。')

h2('4.3 个性化视频/资源推荐')
para('系统按学生对话中的主题检索学习视频，结合主题相关性与播放量热度排序，并通过教学信号词与娱乐/游戏负向词过滤，'
     '只保留权威学习类内容。当真实检索接口不可用时，降级为指向中国大学MOOC、B站、学堂在线的多平台搜索直达卡片，'
     '保证学生索要资源时始终有可点击的权威入口。')

h2('4.4 防幻觉与内容安全')
para('独立的 ContentSafetyService 提供三类能力：①敏感/违规词检测、输入拦截与输出脱敏；'
     '②可注入 system prompt 的防幻觉约束（不编造、不确定显式说明、紧扣知识库）；③学术事实性约束。'
     '已接入对话主流程：违规输入直接拒绝不进入大模型，所有对话的系统提示词末尾统一追加可靠性与安全约束。')

h2('4.5 大模型适配与高可用回退')
para('LlmService 作为统一适配/路由层，按配置优先使用科大讯飞星火；当主用未配置或调用失败（且尚未吐出 token）时，'
     '自动回退到 DeepSeek，保证多智能体稳定可用。另提供离线演示模式，在无网络/无密钥环境下返回本地构造响应，便于比赛演示。')
print("impl details done")

# ============ 5. 开发中遇到的问题与解决方法 ============
h1('五、开发过程中遇到的问题与解决方法')
para('以下问题均为本项目开发过程中真实出现并已解决，记录现象、根因与解决方案，供复盘与文档佐证。')

def problem(idx, title, phenom, cause, fix):
    h2(f'5.{idx} {title}')
    para('现象：', bold=True); para(phenom)
    para('根因：', bold=True); para(cause)
    para('解决方法：', bold=True); para(fix)

problem(1, 'SSE 流式输出中 Markdown 标题不渲染、换行丢失',
        '前端 AI 回复里加粗能正常显示，但 ## / ### 标题不生效，段落全部黏连在一起。',
        'SSE 协议中换行是事件分隔符。后端把每个 token 当纯文本写入 data 字段，当 token 内容为换行符时直接破坏了 SSE 帧；'
        '前端按空行切分事件，换行 token 被切碎为空串后丢弃，导致依赖独占整行的 Markdown 标题无法被解析。',
        '后端将 message token 改为 JSON 编码（{"content": token}）发送，JSON 会把换行转义为字面字符，wire 上不再有真实换行；'
        '前端统一以 JSON 解析所有事件并还原 content，换行得以完整保留。')

problem(2, 'B 站视频检索被风控降级，导致"推荐视频"无结果',
        '学生在对话中索要视频时，AI 给出文字但没有任何视频卡片，表现为"没真正推荐"。',
        '实测 B 站搜索接口对服务器出口 IP 返回 code=-1200（被降级过滤），携带正确 buvid3 Cookie 仍被拒；'
        '属于基于 IP 的服务端风控，改代码无法绕过。原逻辑在检索为空时静默返回空列表，前端无卡片可显示。',
        '改为多平台降级策略：真实检索为空或异常时，回退生成指向中国大学MOOC、B站、学堂在线的"搜索直达卡片"'
        '（搜索结果页在用户本机浏览器可正常打开）。保证学生索要资源时始终有可点击的权威学习入口。')

problem(3, '视频推荐"有概率"不出现',
        '同样的"再推荐几个"指令，有时出卡片有时不出，行为不稳定。',
        '主题解析链路中唯一会随机失败的环节是 AI 主题抽取（chatJsonSync）。当消息为纯指令、会话未绑定知识点、'
        '且该次 AI 抽取恰好返回空时，主题解析返回 null，根本不进入检索，兜底也无从触发。',
        '在主题解析末尾增加不依赖 AI 的确定性兜底：从最近一轮对话中挖掘上一个有效主题（指令型消息的主题本就在前文），'
        '消除对 AI 抽取的概率依赖。')

problem(4, '推荐的视频卡片与当前主题无关（上下文关联弱）',
        '切换话题后说"推荐视频"，卡片仍跟着几轮前的旧主题走。',
        '上下文关联过弱：历史上下文回看窗口过大且每条截断到 80 字，正在讨论的主题词被截掉；'
        '兜底从 6 轮历史里"找第一个有效主题"，缺乏就近权重，容易命中旧主题。',
        '三处协同收敛：①上下文只取最近 2 轮并放宽截断、对 AI 回复优先保留标题；②兜底窗口收窄到最近 1 轮、'
        '优先取上一条 AI 回复标题；③在 AI 抽取提示词中加入"就近优先、以最近一轮为准"约束。')
print("problems done")

# ============ 6. 开发中必然遇到的难题 ============
h1('六、开发过程中必然遇到的难题与应对')
para('以下是本类"大模型 + 多智能体"系统开发中具有普遍性、几乎必然遇到的工程难题，以及本项目采用的应对策略。')

def challenge(idx, title, desc, strategy):
    h2(f'6.{idx} {title}')
    para('难题：', bold=True); para(desc)
    para('应对：', bold=True); para(strategy)

challenge(1, '大模型输出的不确定性与格式不稳定',
          '大模型对同一输入的输出存在随机性，要求返回严格 JSON 时仍可能夹带解释性文字、Markdown 代码围栏或字段缺失，'
          '导致后端解析失败、功能时好时坏。',
          '①提示词中强制"只返回严格 JSON"并给出字段示例；②后端对返回做容错解析（剥离围栏、宽松反序列化）；'
          '③关键链路设确定性兜底（如主题解析在 AI 失败时回退启发式），不把稳定性单独押在模型上。')

challenge(2, '大模型与外部接口的高延迟与超时',
          '大模型推理与多模态/资源检索普遍耗时数秒甚至更久，同步等待会造成前端长时间白屏，体验差且易触发网关超时。',
          '①对话采用 SSE 流式边生成边渲染；②长连接加心跳保活防代理断连；③资源生成提供进度/流式呈现；'
          '④为 Agent 生成设置超时与快速兜底内容，超时不阻塞主流程。')

challenge(3, '大模型服务的可用性与成本',
          '单一大模型提供商存在限流、密钥失效、网络抖动等风险，比赛演示环境尤其脆弱；调用频繁也带来成本压力。',
          '①统一适配层 + 多提供商自动回退（星火主用、DeepSeek 备用）；②画像等结果 Redis 缓存减少重复调用；'
          '③提供离线演示模式，无网络/无密钥也能完整演示。')

challenge(4, '多智能体协同的编排与防失控',
          '多个智能体相互触发容易形成复杂依赖甚至死循环，且某个 Agent 失败可能拖垮整条链路，调试定位困难。',
          '①事件驱动编排，职责单一、松耦合；②设最大协作轮次上限防级联死循环；'
          '③Pipeline 中单 Agent 失败即中断并保留已完成状态；④记录 AgentRunLog 便于回溯。')

challenge(5, '上下文管理与指代消解',
          '多轮对话中"再来几个""这个怎么理解"等指代，依赖正确的上下文窗口；窗口过大引入旧主题噪声，过小又丢失关键信息。',
          '收敛上下文窗口到最近一两轮、对长回复优先保留标题、并在提示词中强调"就近优先、以最近一轮为准"，'
          '平衡信息完整性与相关性。')

challenge(6, '防幻觉与内容安全',
          '大模型可能编造不存在的公式、文献或结论，也可能被诱导输出不当内容，这在教育场景中风险尤高。',
          '统一的 ContentSafetyService：输入敏感词拦截、输出脱敏、系统提示词注入防幻觉与学术准确性约束，'
          '要求不确定内容显式声明、紧扣给定知识库。')

challenge(7, '跨平台与环境一致性',
          '团队在 Windows/macOS/Linux 混合环境开发，路径分隔符、编码、数据库差异易导致"在我机器上能跑"。',
          '①生产 MySQL + 演示 H2 双数据源配置；②统一 UTF-8 与构建配置；③依赖与版本通过 Maven/npm 锁定，降低环境漂移。')
print("challenges done")

# ============ 7. 结语 ============
h1('七、结语')
para('本系统围绕赛题"个性化资源生成 + 学习多智能体"的核心命题，完成了对话式多维画像、多智能体协同资源生成、'
     '个性化路径规划与推送、智能辅导、学习效果评估等功能，并在流式体验、防幻觉与内容安全、'
     '大模型高可用等非功能性需求上做了系统性工程实践。后续将重点增强多模态资源（图像/语音/动画）的真实生成能力，'
     '进一步提升系统的创新性与实用价值。')

doc.save(r'D:/Idea/中国软件杯/SmartMentor技术实现说明书.docx')
print("SAVED")







